from win32com import client as wc
import os
import docx


# 将.doc文件另存为.docx文件
def save_as(file_name):
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(u'G:\python后续阶段\第三阶段\项目\数据清洗\word\\' + file_name)
    doc.SaveAs(u'G:\python后续阶段\第三阶段\项目\数据清洗\word1\\' + file_name + 'x', 12, False, "", True, "", False, False, False, False)
    doc.Close()
    word.Quit()


# 获取目录下所有的文件名
def get_file_names(file_dir):
    for files in os.walk(file_dir):
        return files


# 解析docx文件中的内容
def get_data_docx(file_name):
    file = docx.Document(file_name)
    # print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段
    #
    # 输出每一段的内容
    # for para in file.tables:
    #     print(para.text)
    # # # 输出段落编号及段落内容
    # for i in range(len(file.paragraphs)):
    #     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
    try:
        data_list = [file.tables[0].rows[2].cells[0].text, file.tables[0].rows[3].cells[0].text, file.tables[0].rows[4].cells[0].text, file.tables[4].rows[0].cells[0].text]
    except:
        return
    # print(data_list)
    name = data_list[0].split('  ')[0]
    sex = data_list[1].split(' | ')[0]
    age = data_list[1].split(' | ')[2]
    if '年' in age:
        age=data_list[1].split(' | ')[1]
    birth = data_list[1].split(' | ')[3]
    if '年' not in birth:
        birth=data_list[1].split(' | ')[2]
    degree = data_list[1].split(' | ')[4]
    if '国' in degree:
        degree=data_list[1].split(' | ')[3]
    country=data_list[1].split(' | ')[5]
    if '国'not in country:
        country = data_list[1].split(' | ')[4]
    info = data_list[2].split('\n')
    phone = email = ''
    for i in info:
        if '联系电话' in i:
            phone = i.split('：')[1]
        elif '电子邮箱' in i:
            email = i.split('：')[1]

    school = data_list[3].split('\u3000\u3000\u3000')[1]
    #
    with open('data.txt', 'a') as f:
        f.write(name + '\t' + sex + '\t' + age + '\t' + birth + '\t' + degree + '\t' + phone + '\t' +
                email + '\t' + school + '\n')
    print(name, sex, age, birth, degree, country,phone, email, school)
    # print(name,sex,age,birth,degree,country)
    # for t in file.tables:
    # #     print('table1111111111111111111')
    # #     for r in t.rows:
    # #         print('row2222222222222222222222222222')
    # #         for c in r.cells:
    # #             # if c.text
    # #             print(c.text)

if __name__ == '__main__':
    # file_names = get_file_names(u'G:\python后续阶段\第三阶段\项目\数据清洗\word')[2]
    # print(file_names)
    # for i in file_names:
    #     save_as(i)
    file_names = get_file_names(u'G:\python后续阶段\第三阶段\项目\数据清洗\word1')
    for i in file_names[2][1:-1]:

        get_data_docx(file_names[0]+'\\'+i)
    # get_data_docx(file_names[0] + '\\' + file_names[2][0])
